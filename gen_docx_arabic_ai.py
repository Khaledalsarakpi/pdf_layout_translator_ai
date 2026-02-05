import cv2
import numpy as np
import fitz
from rapidocr import RapidOCR
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml import parse_xml
from xml.sax.saxutils import escape
import os
import wordninja
import re
import time
from transformers import MarianMTModel, MarianTokenizer
import torch

# Initialize Translation Model (Lazy Load)
model = None
tokenizer = None

def load_model():
    global model, tokenizer
    if model is None:
        print("Loading translation model (Helsinki-NLP/opus-mt-en-ar)...")
        model_name = "Helsinki-NLP/opus-mt-en-ar"
        tokenizer = MarianTokenizer.from_pretrained(model_name)
        model = MarianMTModel.from_pretrained(model_name)

def translate_text(text):
    if not text.strip():
        return ""
    load_model()
    try:
        # Split into sentences or chunks if too long?
        # MarianMT has limit of 512 tokens.
        # Our blocks are usually paragraphs.
        inputs = tokenizer(text, return_tensors="pt", padding=True, truncation=True, max_length=512)
        with torch.no_grad():
            translated = model.generate(**inputs)
        return tokenizer.decode(translated[0], skip_special_tokens=True)
    except Exception as e:
        print(f"Translation Error: {e}")
        return text

# --- XML Helper Functions for Floating Elements ---

def add_float_picture(p, image_path_or_stream, width=None, height=None, pos_x=0, pos_y=0):
    """
    Insert a floating picture.
    pos_x, pos_y in EMU (English Metric Units).
    """
    run = p.add_run()
    # Add picture normally first
    inline_shape = run.add_picture(image_path_or_stream, width=width, height=height)
    
    # Now convert inline to floating
    inline = inline_shape._inline
    
    # Create anchor element
    anchor = OxmlElement('wp:anchor')
    anchor.set('distT', "0")
    anchor.set('distB', "0")
    anchor.set('distL', "0")
    anchor.set('distR', "0")
    anchor.set('simplePos', "0")
    anchor.set('relativeHeight', "251658240")
    anchor.set('behindDoc', "1") # Send to back to avoid covering text
    anchor.set('locked', "0")
    anchor.set('layoutInCell', "1")
    anchor.set('allowOverlap', "1")
    
    # SimplePos
    simplePos = OxmlElement('wp:simplePos')
    simplePos.set('x', "0")
    simplePos.set('y', "0")
    anchor.append(simplePos)
    
    # Position H
    positionH = OxmlElement('wp:positionH')
    positionH.set('relativeFrom', "page")
    posOffsetH = OxmlElement('wp:posOffset')
    posOffsetH.text = str(int(pos_x))
    positionH.append(posOffsetH)
    anchor.append(positionH)
    
    # Position V
    positionV = OxmlElement('wp:positionV')
    positionV.set('relativeFrom', "page")
    posOffsetV = OxmlElement('wp:posOffset')
    posOffsetV.text = str(int(pos_y))
    positionV.append(posOffsetV)
    anchor.append(positionV)
    
    # Extent (Size)
    extent = OxmlElement('wp:extent')
    extent.set('cx', str(inline.extent.cx))
    extent.set('cy', str(inline.extent.cy))
    anchor.append(extent)
    
    # Wrap None
    wrapNone = OxmlElement('wp:wrapNone')
    anchor.append(wrapNone)
    
    # DocPr
    docPr = OxmlElement('wp:docPr')
    docPr.set('id', '666')
    docPr.set('name', 'Image')
    anchor.append(docPr)
    
    # Graphic
    graphic = inline.graphic
    anchor.append(graphic)
    
    # Replace inline with anchor
    parent = inline.getparent()
    parent.replace(inline, anchor)
    
    return inline_shape

def create_textbox(p, text, x_pt, y_pt, w_pt, h_pt, fontsize_pt=10, is_rtl=False):
    """
    Create a floating textbox at absolute position.
    Coordinates in points.
    Anchored to the provided paragraph 'p'.
    """
    run = p.add_run()
    
    # Escape text for XML
    text = escape(text)
    
    # VML Injection with correct namespaces and styling
    # Must be wrapped in w:pict for valid Word XML
    
    # RTL properties
    pPr_xml = ""
    rPr_extra = ""
    textbox_style_extra = ""
    
    if is_rtl:
        # Paragraph properties: RTL direction, Justified Alignment
        pPr_xml = '<w:pPr><w:bidi w:val="1"/><w:jc w:val="both"/><w:rPr><w:lang w:bidi="ar-SA"/></w:rPr></w:pPr>'
        
        # Run properties: RTL, Arabic Font
        rPr_extra = '<w:rtl w:val="1"/><w:lang w:bidi="ar-SA"/><w:rFonts w:ascii="Times New Roman" w:hAnsi="Times New Roman" w:cs="Times New Roman" w:hint="cs"/>'
        
        # Textbox style: backup direction
        textbox_style_extra = ";direction:rtl;text-align:justify"

    vml_xml = f"""
    <w:pict xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" 
            xmlns:v="urn:schemas-microsoft-com:vml">
        <v:rect style="position:absolute;margin-left:{x_pt}pt;margin-top:{y_pt}pt;width:{w_pt}pt;height:{h_pt}pt;z-index:251659264;v-text-anchor:top;mso-position-vertical-relative:page;mso-position-horizontal-relative:page"
                filled="false" stroked="false">
          <v:textbox style="mso-fit-shape-to-text:true{textbox_style_extra}">
            <w:txbxContent>
              <w:p>
                {pPr_xml}
                <w:r>
                  <w:rPr>
                    <w:sz w:val="{int(fontsize_pt * 2)}"/>
                    <w:color w:val="000000"/>
                    {rPr_extra}
                  </w:rPr>
                  <w:t>{text}</w:t>
                </w:r>
              </w:p>
            </w:txbxContent>
          </v:textbox>
        </v:rect>
    </w:pict>
    """
    
    pict = parse_xml(vml_xml)
    run._r.append(pict)

# --- Helper for Math Detection ---
def is_likely_math(text):
    # If text contains typical math operators or very few letters
    # Removed + and - as they are common in text
    math_symbols = set("=<>×÷∫∑√≤≥≈≠±∞")
    
    # Check for math symbols (must have at least one strong indicator)
    if any(s in text for s in math_symbols):
        return True, "Symbol Match"
        
    # Check for math-like patterns: "x + y", "f(x)", "A = B"
    # Regex for variable-operator-variable (stricter)
    if re.search(r'[a-zA-Z0-9]\s*[\=\<\>]\s*[a-zA-Z0-9]', text):
        return True, "Regex Pattern"

    # Check for single letters or variables like "x", "y" (too aggressive maybe?)
    # Instead, check ratio of non-alpha characters
    alpha_count = sum(c.isalpha() for c in text)
    # If text is short and mostly symbols/digits
    if len(text) > 0 and (alpha_count / len(text) < 0.4): # Lowered threshold to 0.4
        # Allow simple figure labels like "Fig. 1" to pass as text
        if "Fig" in text or "Figure" in text:
            return False, "Figure Caption"
        return True, "Low Alpha Ratio"
        
    # Check for patterns like "Fig 10.1" (keep this as text, NOT math)
    if re.search(r'Fig\.|Figure', text, re.IGNORECASE):
        return False, "Figure Caption"
        
    return False, "Text"

# --- Main Logic ---

def process_all_pages(pdf_file=None, docx_file=None, max_pages=None):
    if pdf_file is None:
        pdf_file = 'Point, Line, and Edge Detection.pdf'
    
    if docx_file is None:
        base_name = os.path.splitext(os.path.basename(pdf_file))[0]
        docx_file = f"{base_name}_Arabic_AI.docx"
    
    doc_pdf = fitz.open(pdf_file)
    
    # We use global 'translate_text' function now
    
    doc_word = Document()
    
    # Set page size (A4 default? We should match PDF)
    # Note: We set section properties based on the first page
    # If pages have different sizes, we should handle it per page/section.
    page_0 = doc_pdf[0]
    section = doc_word.sections[0]
    section.page_width = Pt(page_0.rect.width)
    section.page_height = Pt(page_0.rect.height)
    section.top_margin = Pt(0)
    section.bottom_margin = Pt(0)
    section.left_margin = Pt(0)
    section.right_margin = Pt(0)
    
    ocr = RapidOCR()
    
    for page_num, page in enumerate(doc_pdf):
        if max_pages is not None and page_num >= max_pages:
            print(f"Reached page limit of {max_pages}. Stopping.")
            break
            
        # Limit removed for full processing
        print(f"Processing page {page_num + 1}/{len(doc_pdf)}...")
        
        # Add page break if not first page
        if page_num > 0:
            doc_word.add_page_break()
            
            # For subsequent pages, we might need a new section if we wanted to change page size
            # But here we assume consistent page size.
        
        # We need a paragraph to anchor elements to for THIS page
        # Adding a paragraph at the current cursor position (which is on the new page)
        p_anchor = doc_word.add_paragraph()
        
        # Page geometry
        page_w = page.rect.width
        page_h = page.rect.height
        
        # OCR and Segmentation
        pix = page.get_pixmap(dpi=300) 
        img_data = pix.tobytes("png")
        
        nparr = np.frombuffer(img_data, np.uint8)
        img_cv = cv2.imdecode(nparr, cv2.IMREAD_COLOR)
        
        result = ocr(img_data)
        
        scale_x = page_w / pix.width
        scale_y = page_h / pix.height
        
        # 1. Extract Figures
        mask_text = np.zeros(img_cv.shape[:2], dtype=np.uint8)
        
        valid_indices = []
        if result and hasattr(result, 'boxes') and hasattr(result, 'scores') and hasattr(result, 'txts'):
            boxes = result.boxes
            scores = result.scores
            txts = result.txts
            if boxes is not None and len(boxes) > 0:
                for i, box in enumerate(boxes):
                    # Filter out low confidence text (likely math or noise)
                    # Threshold 0.70 is more aggressive to catch math
                    if scores[i] < 0.70: 
                        continue
                    
                    # Filter out short text or symbol-heavy text (likely math garbage)
                    txt_val = txts[i].strip()
                    if len(txt_val) < 3:
                        continue
                        
                    num_alnum = sum(c.isalnum() for c in txt_val)
                    if len(txt_val) > 0 and (num_alnum / len(txt_val) < 0.4):
                        continue
                        
                    valid_indices.append(i)
                    cnt = np.array(box).astype(np.int32)
                    cv2.fillPoly(mask_text, [cnt], 255)
                
        kernel_text = np.ones((10,10), np.uint8)
        mask_text_dilated = cv2.dilate(mask_text, kernel_text, iterations=1)
        
        img_no_text = img_cv.copy()
        img_no_text[mask_text_dilated == 255] = [255, 255, 255]
        
        gray = cv2.cvtColor(img_no_text, cv2.COLOR_BGR2GRAY)
        _, thresh = cv2.threshold(gray, 240, 255, cv2.THRESH_BINARY_INV)
        kernel_fig = np.ones((20,20), np.uint8)
        thresh_dilated = cv2.dilate(thresh, kernel_fig, iterations=2)
        contours, _ = cv2.findContours(thresh_dilated, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
        
        print(f"  Page {page_num + 1}: Adding {len(contours)} figures...")
        
        min_area = 1000 # Lowered to capture small equations
        for i, cnt in enumerate(contours):
            x, y, w, h = cv2.boundingRect(cnt)
            if w * h > min_area:
                roi = img_cv[y:y+h, x:x+w]
                
                temp_img_path = f"temp_fig_p{page_num}_{i}.png"
                cv2.imwrite(temp_img_path, roi)
                
                pos_x_pt = x * scale_x
                pos_y_pt = y * scale_y
                
                width_emu = int(w * scale_x * 12700)
                height_emu = int(h * scale_y * 12700)
                
                pos_x_emu = int(pos_x_pt * 12700)
                pos_y_emu = int(pos_y_pt * 12700)
                
                try:
                    add_float_picture(p_anchor, temp_img_path, width=width_emu, height=height_emu, pos_x=pos_x_emu, pos_y=pos_y_emu)
                except Exception as e:
                    print(f"Error adding picture: {e}")
                
                if os.path.exists(temp_img_path):
                    os.remove(temp_img_path)

        # 2. Add Textboxes
        print(f"  Page {page_num + 1}: Adding textboxes...")
        if result and hasattr(result, 'boxes') and hasattr(result, 'txts') and hasattr(result, 'scores'):
            boxes = result.boxes
            txts = result.txts
            scores = result.scores
            
            if boxes is not None:
                text_items = []
                
                # Only process valid indices (high confidence text)
                for i in valid_indices:
                    text = txts[i]
                    box = boxes[i]
                    
                    # --- Text Cleaning Logic ---
                    tokens = text.split()
                    new_tokens = []
                    for token in tokens:
                        if len(token) > 4 and not re.match(r'^[0-9\.\-\+\(\)]+$', token):
                            token_fixed = re.sub(r'([a-z])([A-Z])', r'\1 \2', token)
                            
                            if len(token_fixed) > 12 and token_fixed.isalpha() and token_fixed.islower():
                                 split_list = wordninja.split(token_fixed)
                                 token_fixed = " ".join(split_list)
                            
                            if len(token_fixed) > 10 and token_fixed.isalpha():
                                 split_list = wordninja.split(token_fixed)
                                 if len(split_list) > 1:
                                     token_fixed = " ".join(split_list)
                            
                            new_tokens.append(token_fixed)
                        else:
                            new_tokens.append(token)
                    
                    text = " ".join(new_tokens)
                    text = re.sub(r'([a-z0-9])([.,])([a-zA-Z])', r'\1\2 \3', text)
                    text = re.sub(r'([a-z])([A-Z])', r'\1 \2', text)

                    def split_long_alpha(match):
                        word = match.group(0)
                        split_words = wordninja.split(word)
                        if len(split_words) > 1:
                            if word[0].isupper():
                                split_words[0] = split_words[0].capitalize()
                            return " ".join(split_words)
                        return word
                    
                    text = re.sub(r'[a-zA-Z]{6,}', split_long_alpha, text)
                    
                    xs = [float(p[0]) * scale_x for p in box]
                    ys = [float(p[1]) * scale_y for p in box]
                    x0, y0 = min(xs), min(ys)
                    x1, y1 = max(xs), max(ys)
                    h_pt = y1 - y0
                    
                    text_items.append({
                        'text': text,
                        'x0': x0,
                        'y0': y0,
                        'x1': x1,
                        'y1': y1,
                        'h': h_pt
                    })

                text_items.sort(key=lambda x: x['y0'])
                
                merged_blocks = []
                if text_items:
                    current_block = [text_items[0]]
                    
                    for item in text_items[1:]:
                        prev = current_block[-1]
                        
                        vertical_gap = item['y0'] - prev['y1']
                        align_diff = abs(item['x0'] - prev['x0'])
                        height_diff = abs(item['h'] - prev['h'])
                        
                        MAX_GAP = prev['h'] * 1.5
                        MAX_ALIGN = 50.0
                        MAX_H_DIFF = 5.0
                        
                        # Check math consistency to avoid merging math with text
                        is_prev_math, _ = is_likely_math(prev['text'])
                        is_curr_math, _ = is_likely_math(item['text'])
                        
                        if vertical_gap < MAX_GAP and align_diff < MAX_ALIGN and height_diff < MAX_H_DIFF and (is_prev_math == is_curr_math):
                            current_block.append(item)
                        else:
                            merged_blocks.append(current_block)
                            current_block = [item]
                    
                    merged_blocks.append(current_block)
                
                print(f"  Page {page_num + 1}: Merged {len(text_items)} lines into {len(merged_blocks)} text blocks.")
                
                for idx, block in enumerate(merged_blocks):
                    b_x0 = min(item['x0'] for item in block)
                    b_y0 = min(item['y0'] for item in block)
                    b_x1 = max(item['x1'] for item in block)
                    b_y1 = max(item['y1'] for item in block)
                    
                    b_w = b_x1 - b_x0
                    b_h = b_y1 - b_y0
                    
                    full_text = " ".join([item['text'] for item in block])
                    
                    avg_h = sum(item['h'] for item in block) / len(block)
                    fontsize = avg_h * 0.7
                    if fontsize < 6: fontsize = 6
                    
                    b_w += 5
                    
                    # Math detection
                    is_math, math_reason = is_likely_math(full_text)
                    if is_math:
                        print(f"    Math detected ({math_reason}). Inserting image clip...")
                        
                        # Convert points back to pixels for cropping
                        x_px = int(b_x0 / scale_x)
                        y_px = int(b_y0 / scale_y)
                        w_px = int(b_w / scale_x)
                        h_px = int(b_h / scale_y)
                        
                        # Ensure bounds
                        y_px = max(0, y_px)
                        x_px = max(0, x_px)
                        y_end = min(img_cv.shape[0], y_px + h_px)
                        x_end = min(img_cv.shape[1], x_px + w_px)
                        
                        if x_end > x_px and y_end > y_px:
                             crop_img = img_cv[y_px:y_end, x_px:x_end]
                             crop_path = f"temp_math_p{page_num}_{idx}.png"
                             cv2.imwrite(crop_path, crop_img)
                             
                             # Convert points to EMU for add_float_picture
                             pos_x_emu = int(b_x0 * 12700)
                             pos_y_emu = int(b_y0 * 12700)
                             width_emu = int(b_w * 12700)
                             height_emu = int(b_h * 12700)
                             
                             try:
                                 add_float_picture(p_anchor, crop_path, width=width_emu, height=height_emu, pos_x=pos_x_emu, pos_y=pos_y_emu)
                             except Exception as e:
                                 print(f"Error adding math picture: {e}")
                             
                             if os.path.exists(crop_path):
                                 os.remove(crop_path)
                    else:
                        # Translate
                        print(f"    Translating: {full_text[:30]}...")
                        translated_text = translate_text(full_text)
                        
                        if translated_text == full_text and len(full_text) > 5:
                             print(f"    [WARN] Translation returned same text: {full_text[:30]}...")
                             create_textbox(p_anchor, full_text, b_x0, b_y0, b_w, b_h, fontsize, is_rtl=False)
                        else:
                             print(f"    Result: {translated_text[:30]}...")
                             create_textbox(p_anchor, translated_text, b_x0, b_y0, b_w, b_h, fontsize, is_rtl=True)
                
                # Force save every 5 pages to avoid total loss
                if (page_num + 1) % 5 == 0:
                     doc_word.save(docx_file)
                     print(f"  [Checkpoint] Saved progress to {docx_file}")

    doc_word.save(docx_file)
    print(f"Saved full docx to {docx_file}")

if __name__ == "__main__":
    process_all_pages()
